import streamlit as st
missing = []
# Defensive imports: collect which modules are missing but don't crash the app
try:
    from PIL import Image
except Exception:
    Image = None
    missing.append("Pillow (PIL)")

try:
    import requests
except Exception:
    requests = None
    missing.append("requests")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    missing.append("python-docx")

try:
    import openai
except Exception:
    openai = None
    missing.append("openai")

try:
    import base64, io, re
except Exception:
    base64 = None
    io = None
    re = None

st.set_page_config(page_title="Debuggable Highlight App", page_icon="🛠️", layout="wide")
st.title("🛠️ Debug helper — Highlight app (defensive mode)")

if missing:
    st.error("האפליקציה לא תוכל לפעול במלואה. ישנם מודולים חסרים:")
    for m in missing:
        st.write(f"- **{m}**")
    st.info("אופציות להמשך:")
    st.write("1. ודא שקובץ `requirements.txt` בפרויקט כולל את השורות הבאות:")
    st.code('''streamlit
pillow
requests
python-docx
openai''')
    st.write("2. בדוק בלשונית *Manage app → Logs* ב-Streamlit Cloud לקבלת שגיאות התקנה מלאות.")
    st.write("3. אם אתה רץ מקומית: הרץ `pip install -r requirements.txt` בספריית הפרוייקט.")
    st.write("4. אם מודול 'pillow' קיים אבל עדיין יש בעיה, נסה להריץ `pip install --upgrade pip setuptools wheel` ואז להתקין שוב.")
    st.markdown("---")

# If essential modules are missing, stop here but keep the app responsive
essentials = [Image, requests, Document, openai]
if any(x is None for x in essentials):
    st.warning("כרגע חסרים מודולים קריטיים ולכן ממשק הפעולה המרכזי מושבת. לאחר שתתקין את החבילות, רענן את האפליקציה.")
    st.write("להלן טקסט מוצע ל-Secrets (Manage app → Secrets):")
    st.code('''OCR_API_KEY="helloworld"
TMDB_API_KEY="your_tmdb_key"
OPENAI_API_KEY="your_openai_key"''')
    st.stop()

# If we reach here, essential modules are present; include the minimal working app (OCR via OCR.space)
OCR_KEY = st.secrets.get("OCR_API_KEY")
TMDB_KEY = st.secrets.get("TMDB_API_KEY")
openai.api_key = st.secrets.get("OPENAI_API_KEY")

def extract_text_from_image(image):
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_base64 = base64.b64encode(buffered.getvalue()).decode()
    payload = {
        "base64Image": "data:image/png;base64," + img_base64,
        "language": "eng,heb",
        "apikey": OCR_KEY or "helloworld",
        "isOverlayRequired": False
    }
    try:
        r = requests.post("https://api.ocr.space/parse/image", data=payload, timeout=30)
        result = r.json()
        return result.get("ParsedResults", [{}])[0].get("ParsedText", "").strip()
    except Exception as e:
        st.error(f"OCR API error: {e}")
        return ""

def clean_text(t):
    return re.sub(r"[^a-zA-Z0-9א-ת ]", "", (t or "").replace("\n", " ").strip())

def search_series_info(series_name):
    if not TMDB_KEY:
        return {"name": series_name, "overview": "TMDB API key חסר", "first_air_date": "לא ידוע", "episodes": "לא ידוע"}
    try:
        url = f"https://api.themoviedb.org/3/search/tv?api_key={TMDB_KEY}&query={series_name}"
        r = requests.get(url, timeout=10)
        data = r.json()
        if not data.get("results"):
            return {"name": series_name, "overview": "לא נמצא מידע", "first_air_date": "לא ידוע", "episodes": "לא ידוע"}
        s = data["results"][0]
        return {
            "name": s.get("name", series_name),
            "overview": s.get("overview", "אין תקציר"),
            "first_air_date": s.get("first_air_date", "לא ידוע"),
            "episodes": s.get("number_of_episodes", "לא ידוע")
        }
    except Exception:
        return {"name": series_name, "overview": "שגיאת API", "first_air_date": "לא ידוע", "episodes": "לא ידוע"}

def generate_summary(text):
    if not openai.api_key:
        return "חסר מפתח OpenAI"
    try:
        r = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "סכם טקסטים בעברית."},
                {"role": "user", "content": f"צור תקציר בעברית:\n{text}"}
            ],
            max_tokens=200,
            temperature=0.2
        )
        return r["choices"][0]["message"]["content"]
    except Exception:
        return "שגיאה ביצירת תקציר"

def create_doc(series_list):
    doc = Document()
    title = doc.add_paragraph("היילייטס סדרות", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for s in series_list:
        doc.add_paragraph(s["name"], style="Heading 1")
        p = doc.add_paragraph(f"תאריך עלייה: {s['first_air_date']}")
        p.runs[0].font.color.rgb = RGBColor(0,0,128)
        doc.add_paragraph(f"מספר פרקים: {s['episodes']}")
        doc.add_paragraph("תקציר:", style="Heading 2")
        doc.add_paragraph(s["summary"])
        doc.add_paragraph("----------------------------------------")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.title("📺 היילייטס סדרות – פועל")
uploaded = st.file_uploader("העלה תמונות", type=["jpg","jpeg","png"], accept_multiple_files=True)
if uploaded:
    series_list = []
    for img in uploaded:
        st.markdown("---")
        col1, col2 = st.columns([1,2])
        with col1:
            st.image(img, caption=img.name, use_column_width=True)
        with col2:
            with st.spinner("מבצע OCR..."):
                text = extract_text_from_image(Image.open(img))
            cleaned = clean_text(text)
            st.write("טקסט שחולץ:", cleaned or "_לא נמצא_")
            if not cleaned:
                st.warning("לא נמצא טקסט.")
                continue
            with st.spinner("מחפש מידע..."):
